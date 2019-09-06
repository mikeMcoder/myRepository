# -*- coding: utf-8 -*-
"""
Created on Sat Sep  3 12:48:53 2016

@author: mark
"""

# note: ipywidgets is not required, but seaborn throws a warning without it

import sys
import time
import os
#import copy

sys.path.append("..") 

import struct
import math
import operator
import serial

import matplotlib
import matplotlib.pyplot as plt

from Backbone.Backbone import Backbone

import numpy as np
import pandas as pd

try:
    import docx
    from docx.shared import Pt
except:
    pass

try:
    import seaborn as sns
except ImportError:
    seabornavailable = False
else:
    seabornavailable = True

# do not share the DoNotShare.py file with customer.  Instead, remove inheritance and provide bridge file

try:
    from DoNotShare.DoNotShare import DoNotShare
except:
    class DoNotShare:
        pass

#from enum import Enum

PYTHON32BIT = sys.maxsize < 2**32

RX_TIMEOUT_DEFAULT  = 0.5 # flash writes are slow (~200ms)

class MZcontrol(Backbone,DoNotShare):
    
    def __init__(self, lsr=None, it=None):
        self._kvDictionaryname = 'kvDictionary'
        self._bridgedir = os.getcwd() + os.sep + 'mzbridge'
        self._cachedir = os.getcwd() + os.sep + 'cache'

        handle = None
        
        if lsr is not None:
            handle = lsr._ttyhandle
        elif it is not None:
            handle = it._link

        super(MZcontrol, self).__init__(_ttyhandle=handle)
        
        self._cmd_id_setbits   = 0x02
        self._cmd_id_clearbits = 0

    class CTRLCONTAINER():
        """ """
        pass


    def writedocumentation(self, docfilename='MZcontrol.docx'):
        """Write MZcontrol documentation to file 

        :param docfilename: 

        """

        print "Writing documentation to", docfilename

        doc = docx.Document()
        
        doc.add_heading('MZcontrol Functions', level=1)

        doc.add_paragraph(self.get_MZcontrol_Version())

        doc.add_heading('Commands', level=2)
        
                
        excluded = [attr for attr in dir(self) if callable(getattr(self, attr)) 
                    and not attr.startswith("_") 
                    and getattr(self, attr).__doc__ is None]
        
        if len(excluded) > 0:
            print self.hilite("Warning: Excluding functions missing doc string:\n" + '\n'.join(excluded), 0,1)
        
        members = [(attr,getattr(self, attr).__doc__.strip()) for attr in dir(self) if callable(getattr(self, attr)) 
                    and not attr.startswith("_") 
                    and getattr(self, attr).__doc__ is not None
                    and len(getattr(self, attr).__doc__.strip()) > 0
                    and not getattr(self, attr).__doc__.startswith('\n')]

        membertable = pd.DataFrame(members, columns=['function','description'])
            
        membertable['function'] = 'mz.' + membertable['function'] + '()'

        t = doc.add_table(membertable.shape[0]+1, membertable.shape[1], style=doc.styles['TableGrid'])
        t.autofit = False
        t.columns[0].width = docx.shared.Inches(1.5)
        t.columns[1].width = docx.shared.Inches(6.0 - 1.5)
        
        # add the header rows.
        for j in range(membertable.shape[-1]):
            t.cell(0,j).text = membertable.columns[j]
        
        # add the rest of the data frame
        for i in range(membertable.shape[0]):
            if not i%5:
                sys.stdout.write('.')
                sys.stdout.flush()
            for j in range(membertable.shape[-1]):
                t.cell(i+1,j).text = str(membertable.values[i,j]).strip()
                for paragraph in t.cell(i+1,j).paragraphs:
                    for run in paragraph.runs:
                        font = run.font
                        font.size = Pt(8)
                        if membertable.columns[j] == 'function':
                            font.size = Pt(7)
                            font.name = 'Courier New'
                
                
        dfcal = self._caltable.copy()
                
        dfcal = dfcal[[u'type', u'variable', u'Units', u'comment', u'remote',u'cal']]
        
        #mcudict = {True: 'Master',
        #               False: 'ABC',
        #               }

        #dfcal['MCU'] = dfcal['remote'].map(mcudict)
        

        #dfcal["Function"] = 'mz.cal.'+dfcal['cal']+'.'+dfcal['variable']+'(setval=None)'
        
        dfcal.comment = 'Description : ' + dfcal.comment
        dfcal.loc[dfcal.Units != '', 'comment'] += '\nUnits : ' + dfcal.loc[dfcal.Units != '', 'Units']
        dfcal.loc[dfcal.type != '', 'comment'] += '\nType : ' + dfcal.loc[dfcal.type != '', 'type']

        dfcal.rename(columns={'comment': 'Info','variable': 'Variable'}, inplace=True)

        del dfcal['remote']
        del dfcal['type']
        del dfcal['Units']
        #del dfcal['cal']
        #del dfcal['variable']

        doc.add_heading('Calibration settings', level=2)

        tables = dfcal.groupby('cal', as_index=False) 
        
        for group in tables.groups:
            table = tables.get_group(group)
            del table['cal']
                        
            doc.add_heading('Struct : ' + group, level=3)
            
            doc.add_paragraph('E.g.')
            p = doc.add_paragraph('mz.cal.'+group+'.'+table.values[0,0]+'()')
            p.runs[0].font.size=Pt(8)
            p.runs[0].font.name = 'Courier New'

            t = doc.add_table(table.shape[0]+1, table.shape[1], style=doc.styles['TableGrid'])
            t.autofit = False
            t.columns[0].width = docx.shared.Inches(2.5)
            t.columns[1].width = docx.shared.Inches(6.0 - 2.5)

            # add the header rows.
            for j in range(table.shape[-1]):
                t.cell(0,j).text = table.columns[j]
            
            # add the rest of the data frame
            for i in range(table.shape[0]):
                if not i%5:
                    sys.stdout.write('.')
                    sys.stdout.flush()
                for j in range(table.shape[-1]):
                    t.cell(i+1,j).text = str(table.values[i,j]).strip()
                    for paragraph in t.cell(i+1,j).paragraphs:
                        for run in paragraph.runs:
                            font = run.font
                            font.size = Pt(8)
                            if table.columns[j] == 'Variable':
                                font.size = Pt(7)
                                font.name = 'Courier New'
        


        #doc.save(docfilename)

        doc.add_page_break()

        df = self._functionstable.copy()
        
        df = df[df['get'] != '' ] # remove missing

        df.loc[df.set == '', "Function"] = 'mz.ctrl.'+df['class']+'.'+df['get']+'()'
        
        df.sort_values(by=['table','index'],ascending=[True, True], inplace=True)

        #df['writable']= df['set'] != ''
        df.loc[df.set == '', "Function"] = 'mz.ctrl.'+df['table']+'.'+df['get']+'()'
        df.loc[df.set != '', "Function"] = 'mz.ctrl.'+df['table']+'.'+df['get']+'(setval=None)'

        df = df[['Function','Description','table','Units','Range min','Range max']]
        
        df.rename(columns={'Range min': 'Min', 'Range max': 'Max'}, inplace=True)
        
        df['Range'] = df['Min'] + '--' + df['Max'] + ' ' + df['Units']
        df.loc[df.Range == '-- ', 'Range'] = ''

        doc.add_heading('Low-level Hardware Control', level=2)

#        lasttable=''
#        i=0
#
#
#        for n,row in df.iterrows():
#            i=i+1
#            if not i%5:
#                sys.stdout.write('.')
#                sys.stdout.flush()
#            if lasttable != row['table']:
#                lasttable = row['table']
#                doc.add_heading(row['table'], level=2)
#            doc.add_heading(row['Function'], level=4)
#            self._addlinetodoc(doc, row['Description'])
#            self._addlinetodoc(doc, row['Range'])
#        
#        doc.add_page_break()
        
        del df['Min']
        del df['Max']
        del df['Units']
        
        tables = df.groupby('table', as_index=False) 
        
        for group in tables.groups:
            table = tables.get_group(group)
            del table['table']
            
            if (table['Range'] == '').all():
                del table['Range']
            
            doc.add_heading(group, level=3)

            doc.add_paragraph('E.g.')
            p = doc.add_paragraph(table.values[0,0])
            p.runs[0].font.size=Pt(8)
            p.runs[0].font.name = 'Courier New'

            t = doc.add_table(table.shape[0]+1, table.shape[1], style=doc.styles['TableGrid'])
            t.autofit = False
            t.columns[0].width = docx.shared.Inches(3.0)
            t.columns[1].width = docx.shared.Inches(6.0 - 3.0 - 0.7)
            if len(t.columns) > 2:
                t.columns[2].width = docx.shared.Inches(0.7)
            # add the header rows.
            for j in range(table.shape[-1]):
                t.cell(0,j).text = table.columns[j]
            
            # add the rest of the data frame
            for i in range(table.shape[0]):
                if not i%5:
                    sys.stdout.write('.')
                    sys.stdout.flush()
                for j in range(table.shape[-1]):
                    t.cell(i+1,j).text = str(table.values[i,j]).strip()
                    for paragraph in t.cell(i+1,j).paragraphs:
                        for run in paragraph.runs:
                            font = run.font
                            font.size = Pt(8)
                            if table.columns[j] == 'Function':
                                font.size = Pt(7)
                                font.name = 'Courier New'
        
        
        
        doc.save(docfilename)
        

    def hwstate(self, rawOpt=False):
        """

        :param rawOpt:  (Default value = False)

        """
        retval = self.formatted() #{}
        for hwname, hw in self.ctrl.__dict__.iteritems():
            retval0 = self.formatted()
            retval0._width = 40
            retval0._columns = 2
            for name, val in sorted(hw.__dict__.iteritems()):
                if callable(val):
                    value = val(rawOpt=rawOpt)
                    setattr(retval0, name, value )
            #retval.update({hwname:retval0})
            setattr(retval, hwname, retval0)

#        if hasattr(retval,'P3V3_ADC'):
#            retval.calc_module_pwr = self._floatwithunits(retval.P3V3_ADC * retval.MOD_I_ADC, units='W')
#        else:
#            retval.calc_module_pwr = self._floatwithunits(retval.ABC_AVDD_ADC * retval.MOD_I_ADC, units='W')

        return retval

    def taskstatus(self):
        """ """
        return self.writereg(0xFD).split('\x00')[0]

    def runtimestats(self):
        """ """
        return self.writereg(0xFE).split('\x00')[0]

    def intpriorities(self):
        """ """
        return self.writereg(0xFC).split('\x00')[0]

    def power(self, level=None):
        """powers the module up(1) or down(0)

        :param level:  (Default value = None)

        """
        'Get/Set debug level'
        if level is not None:
            return self.reg(0x38, level, write=True)
        else:
            return self.reg(0x38)
    

    def lsrOn(self, LaserOn = 0):
        pos = 0
        flags = 0x0B      #Remote
        regnum       = 0xDB                               
        mainCtrlDbg  = 9
        laserCommand = 3
        returndat = self.sendcommand(chr(flags) + chr(regnum) + chr(mainCtrlDbg) + 
             chr(laserCommand) + chr(LaserOn) + chr(0) + chr(0) + chr(0) )

        rawdata=returndat[pos:pos+2]
        decodeddata = hex(struct.unpack('>' + "H",rawdata)[0])
        print "ResEnab: ",decodeddata
            
    def mainCtrlCmd(self, remote=1, command=0, index=0, charval1=0, charval2=0,
               charval3=0, charval4=0):             # Command: 1=FLASH control, 0=Power control
        regnum = 0xDB                               # FLASH Index: 1=Save Reg Defaults, 2=Save calibration, 3=Invalidate Calib,
        if remote:                                  #    4=Invalidate Defaults
            flags = 0x0B                            # Power Index: 1=Power Up, 0=Power down
        else:
            flags = 0x03
        mainCtrlDbg  = 9
        return self.sendcommand(chr(flags) + chr(regnum) + chr(mainCtrlDbg) +
                                chr(command) + chr(index) +
                                chr(0) + chr(0) +  chr(0))

    def mainCtrlPowerCmd(self, remote=1, index=0):  # Power Index: 0=Power Up, 1=Power down
        if remote:                                  
            flags = 0x0B                            
        else:
            flags = 0x03
        regnum       = 0xDB                               
        mainCtrlDbg  = 9
        PowerCommand = 0
        return self.sendcommand(chr(flags) + chr(regnum) + chr(mainCtrlDbg) +
                                chr(PowerCommand) + chr(index) +
                                chr(0) + chr(0) +  chr(0))

    def mainCtrlFlashCmd(self, remote=1, index=0):  # FLASH Index: 1=Save Reg Defaults (remote only), 2=Save calibration,
        regnum = 0xDB                               #     3=Invalidate Calib, 4=Invalidate Defaults (remote only)
        if remote:                                  
            flags = 0x0B                            
        else:
            flags = 0x03
        mainCtrlDbg  = 9
        flashCommand = 1
        return self.sendcommand(chr(flags) + chr(regnum) + chr(mainCtrlDbg) +
                                chr(flashCommand) + chr(index) +
                                chr(0) + chr(0) +  chr(0))

    def mainCtrlEventCmd(self, remote=1, index=0, arg1=0, arg2=0):  # 0 = systemReset, 1 = systemInitialize
        regnum = 0xDB                               # 2 = systemFault -arg1 = faultsource, arg2 = faultflag
        if remote:                                  # 3 = systemFaultClear -arg1 = faultsource, arg2 = faultflag
            flags = 0x0B                            # 4 = systemLow Power -arg1 = bool assert low power
        else:                                       # 5 = systemTxDisable -arg1 = bool assert tx disable
            flags = 0x03                            # 6 = sysReset&TxDisabDeassert
        mainCtrlDbg  = 9
        EventCommand = 2
        return self.sendcommand(chr(flags) + chr(regnum) + chr(mainCtrlDbg) +
                                chr(EventCommand) + chr(index) +
                                chr(arg1) + chr(arg2) +  chr(0))
                                
    def mainCtrlLaserCmd(self, remote=1, index=0):  # LaserControlCommand Index: 0=laserOff, 1=laserOn
        if remote:                                  # 
            flags = 0x0B                            # 
        else:                                       # 
            flags = 0x03                            # 
        regnum       = 0xDB                         # 
        mainCtrlDbg  = 9
        laserCommand = 3
        return self.sendcommand(chr(flags) + chr(regnum) + chr(mainCtrlDbg) +
                                chr(laserCommand) + chr(index) +
                                chr(0) + chr(0) +  chr(0))
                                
    @staticmethod
    def decodeCriticalFaults(faultList):
        faultDefStructure = [
            "FLASH Faults Master: ",
            "FLASH Faults ABC: ",
            "Laser Faults Master: ",
            "EDFA Faults ABC: ",
            "uMOD Faults ABC: ",
            "Power Faults ABC: ",
            ]
        pos = 0
        for varname in faultDefStructure:
            rawdata=faultList[pos:pos+2]
            decodeddata = hex(struct.unpack('<' + "H",rawdata)[0])
#            decodeFaultFlags(pos, decodeddata, rawdata)
            if(pos==0):
                faultMap = ["REGISTER_DEFAULTS_CRC_ERR, ","REGISTER_VERSION_MISMATCH_ERR, ","REGISTER_FLASH_SAVE_ERR, ","MAIN_CALIB_CRC_ERR, ",
                            "MAIN_CALIB_VERSION_ERR, ","MAIN_CALIB_FLASH_SAVE_ERR, ","UNDEF, ","UNDEF, ",
                            "UNDEF, ","UNDEF, ","UNDEF, ","UNDEF, ","UNDEF, ","UNDEF, ","UNDEF, ","UNDEF, "]
            elif(pos==2):
                faultMap = ["ABC_CALIB_CRC_ERR, ","ABC_CALIB_VERSION_ERR, ","ABC_CALIB_FLASH_SAVE_ERR, ","UNDEF, ", "UNDEF, ","UNDEF, ","UNDEF, ","UNDEF, ",
                            "UNDEF, ","UNDEF, ","UNDEF, ","UNDEF, ","UNDEF, ","UNDEF, ","UNDEF, ","UNDEF, "]
            elif(pos==4):
                faultMap = ["LASER_NOT_PRESENT_ERR, ", "LASER_IN_RESET, ", "LASER_FATAL_ERR, ","LASER_WARNING_ERR, ","LASER_FAILED_TO_TUNE_ERR, ","LASER_TUNED_ON, ","UNDEF, ","UNDEF, ",
                            "UNDEF, ","UNDEF, ","UNDEF, ","UNDEF, ","UNDEF, ","UNDEF, ","UNDEF, ","UNDEF, "]
            elif (pos==6):
                faultMap = ["EDFA_NOT_DETECTED_ERR, ","UNDEF, ","UNDEF, ","UNDEF, ","UNDEF, ","UNDEF, ","UNDEF, ","UNDEF, ",
                            "UNDEF, ","UNDEF, ","UNDEF, ","UNDEF, ","UNDEF, ","UNDEF, ","UNDEF, ","UNDEF, "]
            elif (pos==8):
                faultMap = ["MZ_NOT_DETECTED_ERR, ","UNDEF, ","UNDEF, ","UNDEF, ","UNDEF, ","UNDEF, ","UNDEF, ","UNDEF, ",
                            "UNDEF, ","UNDEF, ","UNDEF, ","UNDEF, ","UNDEF, ","UNDEF, ","UNDEF, ","UNDEF, "]
            elif (pos==10):
                faultMap = ["POWER_ENAB_ERR, ","POWER_5V_ENAB_ERR, ","POWER_P4_12V_X_ENAB_ERR, ","POWER_P4_12V_Y_ENAB_ERR, ","POWER_3V3_DRV_ENAB_ERR, ",
                            "POWER_5V_DRV_ENAB_ERR, ","POWER_UMOD_TEC_ENAB_ERR, ","POWER_P9V_L_ENAB_ERR, ",
                            "UNDEF, ","UNDEF, ","UNDEF, ","UNDEF, ","UNDEF, ","UNDEF, ","UNDEF, ","UNDEF, "]
            else:
                decodeddata += "Unsupported Fault Type!"
                return
            
            bitMask = 0x0001
            faultString = " "
            if(int(decodeddata,16) != 0):
                for bitDefinitions in faultMap:
                    if(int(decodeddata,16) & bitMask):
                        faultString += bitDefinitions      # append bitDefinitions string to decodeddata
                    bitMask = bitMask<<1
            else:                
                decodeddata += " -No Faults"
            pos=pos+2
            print str(varname),decodeddata,str(faultString)
        return

    def getSystemFaults(self, remote=0):   # get all faults.
        if remote:                                  # 
            flags = 0x0B                            # 
        else:                                       # 
            flags = 0x03                            # 
        regnum         = 0xDB                       # 
        getsysfaultCmd = 10
        returndat = self.sendcommand(chr(flags) + chr(regnum) + chr(getsysfaultCmd) +
                                chr(0) + chr(0) +
                                chr(0) + chr(0) +  chr(0))
        self.decodeCriticalFaults(returndat)
        return

    @staticmethod
    def decodeCriticalFaultLists(criticalFaultList):
        print str(criticalFaultList)
#        stateDefStructure = [
#            "MSA_ANY_STATE: ",      #= 100,
#            "MSA_RESET: ",          #= 101, // Low-Power,TX-Off
#            "MSA_INITIALIZE: ",     #= 102, // Low-Power,TX-Off
#            "MSA_LOW_POWER: ",      #= 103, // TX-Off
#            "MSA_HIGH_POWER_UP: ",  #= 104, // TX-Off
#            "MSA_HIGH_POWER_DOWN: ",#= 105, // 
#            "MSA_TX_OFF: ",         #= 106, // 
#            "MSA_TX_TURN_OFF: ",    #= 107, // 
#            "MSA_TX_TURN_ON: ",     #= 108, // 
#            "MSA_READY: ",          #= 109, // 
#            "MSA_FAULT: ",          #= 110  // 
#            ]
        ###############################################################################################################
        # Index into stateDefStructure[state-100]  (see systemCriticalFaultListInit[] in FW) for message structure.
        # Decode just the state field and print it.  Move index past this point in string.  Call decodeCriticalFaults()
        # Move index past this fault record in string.
        # Repeat decode step until record exhausted.
        ###############################################################################################################
        return

    def getCritcalFaults(self):   # get all faults.
        flags = 0x0B      #Remote
        regnum = 0xDB                               # 
        command = 11      # 11 -crititcal fault lists.
        getFaultsCmd = 0
        returndat = self.sendcommand(chr(flags) + chr(regnum) + chr(command) + chr(getFaultsCmd) + chr(0) + chr(0) + chr(0) +  chr(0))
#        self.decodeCriticalFaultLists(returndat)
        return returndat
        
    def critcalFaultsSet(self, state=0, faultType=0, faultFlag=0):   # Set critical fault flags for  a given fault type.
        flags = 0x0B      #Remote
        regnum = 0xDB                               # 
        command = 11      # 11 -crititcal fault lists.
        setFaultCmd = 1
        returndat = self.sendcommand(chr(flags) + chr(regnum) + chr(command) + chr(setFaultCmd) + chr(state)  + chr(faultType) + chr(faultFlag)+  chr(0))
#        self.decodeCriticalFaultLists(returndat)
        return returndat

    def critcalFaultsClearAll(self):   # Clear all critical faults from all lists.
        flags = 0x0B      #Remote
        regnum = 0xDB                               # 
        command = 11      # 11 -crititcal fault lists.
        clrFaultsCmd = 2
        returndat = self.sendcommand(chr(flags) + chr(regnum) + chr(command) + chr(clrFaultsCmd) + chr(0) + chr(0) + chr(0) +  chr(0))
#        self.decodeCriticalFaultLists(returndat)
        return returndat
        
    def critcalFaultsRestoreAll(self):   # Restore all critical faults for all state/fault-type lists.
        flags = 0x0B      #Remote
        regnum = 0xDB                               # 
        command = 11      # 11 -crititcal fault lists.
        clrFaultsCmd = 4
        returndat = self.sendcommand(chr(flags) + chr(regnum) + chr(command) + chr(clrFaultsCmd) + chr(0) + chr(0) + chr(0) +  chr(0))
#        self.decodeCriticalFaultLists(returndat)
        return returndat

    def critcalFaultsClear(self, state=0, faultType=0, faultFlag=0):   # Remove critical fault flags for a given fault type.
        flags = 0x0B      #Remote
        regnum = 0xDB                               # 
        command = 11      # 11 -crititcal fault lists.
        clrFaultCmd = 3
        returndat = self.sendcommand(chr(flags) + chr(regnum) + chr(command) + chr(clrFaultCmd) + chr(state) + chr(faultType) + chr(faultFlag) +  chr(0))
 #       self.decodeCriticalFaultLists(returndat)
        return returndat
    
    def readx99(self, x99version=0):
        """Reports the current MZ controller status"""
        returnblank = False
        returndat = self.writereg(regnum=0x99,index=x99version)
        dataformat = ord(returndat[0])

        defaultdatastructure = [
            ["L","statusbits"           ,"" ],
            ["L","pendingbits"          ,"" ],
            ["f","DCphotodiode"         ,"V"],
            ["f","demodreal1"           ,"" ],
            ["f","demodreal2"           ,"" ],
            ["f","demodreal3"           ,"" ],
            ["f","demodimag1"           ,"" ],
            ["f","demodimag2"           ,"" ],
            ["f","demodimag3"           ,"" ],
            ["f","biasvoltageXI"        ,"V" ],
            ["f","biasvoltageXQ"        ,"V" ],
            ["f","biasvoltageYI"        ,"V" ],
            ["f","biasvoltageYQ"        ,"V" ],
            ["f","biasvoltageX"         ,"V" ],
            ["f","biasvoltageY"         ,"V" ],
            ["f","debugfloat"           ,"V" ],
            ["L","tmpdebugU32"          ,"" ]
        ]
        
        if dataformat == 1:
            datastructure = defaultdatastructure
        elif dataformat == 2:
            datastructure = [
            ["L","statusbits"           ,"" ],
            ["L","pendingbits"          ,"" ],
            ["f","DCphotodiode"         ,""],
            ["f","demodreal1"           ,"" ],
            ["f","demodreal2"           ,"" ],
            ["f","demodreal3"           ,"" ],
            ["f","demodimag1"           ,"" ],
            ["f","demodimag2"           ,"" ],
            ["f","demodimag3"           ,"" ],
            ["f","demod2real1"           ,"" ],
            ["f","demod2real2"           ,"" ],
            ["f","demod2real3"           ,"" ],
            ["f","demod2imag1"           ,"" ],
            ["f","demod2imag2"           ,"" ],
            ["f","demod2imag3"           ,"" ],
            ["f","biasvoltageXI"        ,"V" ],
            ["f","biasvoltageXQ"        ,"V" ],
            ["f","biasvoltageYI"        ,"V" ],
            ["f","biasvoltageYQ"        ,"V" ],
            ["f","biasvoltageX"         ,"V" ],
            ["f","biasvoltageY"         ,"V" ],
            ["f","debugfloat"           ,"V" ],
            ["L","tmpdebugU32"          ,"" ]
        ]
        elif dataformat == 3:
            datastructure = [
            ["L","statusbits"           ,"" ],
            ["L","pendingbits"          ,"" ],
            ["f","DCphotodiode"         ,""],
            ["L","uint8s"   ,""],
            ["f","demodreal1"           ,"" ],
            ["f","demodreal2"           ,"" ],
            ["f","demodreal3"           ,"" ],
            ["f","demodimag1"           ,"" ],
            ["f","demodimag2"           ,"" ],
            ["f","demodimag3"           ,"" ],
            ["f","demod2real1"           ,"" ],
            ["f","demod2real2"           ,"" ],
            ["f","demod2real3"           ,"" ],
            ["f","demod2imag1"           ,"" ],
            ["f","demod2imag2"           ,"" ],
            ["f","demod2imag3"           ,"" ],
            ["f","biasvoltageXI"        ,"V" ],
            ["f","biasvoltageXQ"        ,"V" ],
            ["f","biasvoltageYI"        ,"V" ],
            ["f","biasvoltageYQ"        ,"V" ],
            ["f","biasvoltageX"         ,"V" ],
            ["f","biasvoltageY"         ,"V" ],
            ["f","debugfloat"           ,"V" ],
            ["L","tmpdebugU32"          ,"" ]
        ]
        elif dataformat == 4:
            datastructure = [
            ["L","statusbits"           ,"" ],
            ["L","pendingbits"          ,"" ],
            ["f","DCphotodiode"         ,""],
            ["L","uint8s"   ,""],
            ["f","biasvoltageXI"        ,"V" ],
            ["f","biasvoltageXQ"        ,"V" ],
            ["f","biasvoltageYI"        ,"V" ],
            ["f","biasvoltageYQ"        ,"V" ],
            ["f","biasvoltageX"         ,"V" ],
            ["f","biasvoltageY"         ,"V" ],
            ["f","biaserrorXI_R"          ,"" ],
            ["f","biaserrorXQ_R"          ,"" ],
            ["f","biaserrorYI_R"          ,"" ],
            ["f","biaserrorYQ_R"          ,"" ],
            ["f","biaserrorX_R"           ,"" ],
            ["f","biaserrorY_R"           ,"" ],
            ["f","biaserrorXI_I"          ,"" ],
            ["f","biaserrorXQ_I"          ,"" ],
            ["f","biaserrorYI_I"          ,"" ],
            ["f","biaserrorYQ_I"          ,"" ],
            ["f","biaserrorX_I"           ,"" ],
            ["f","biaserrorY_I"           ,"" ],
            ["f","debugfloat"           ,"V" ],
            ["L","tmpdebugU32"          ,"" ]
        ]
        elif dataformat == 5:
            datastructure = [
            ["L","statusbits"           ,"" ],
            ["L","pendingbits"          ,"" ],
            ["f","DCphotodiode"         ,""],
            ["L","uint8s"   ,""],
            ["f","biasvoltageXI"        ,"V" ],
            ["f","biasvoltageXQ"        ,"V" ],
            ["f","biasvoltageYI"        ,"V" ],
            ["f","biasvoltageYQ"        ,"V" ],
            ["f","biasvoltageX"         ,"V" ],
            ["f","biasvoltageY"         ,"V" ],
            ["f","biaserrorXI"          ,"" ],
            ["f","biaserrorXQ"          ,"" ],
            ["f","biaserrorYI"          ,"" ],
            ["f","biaserrorYQ"          ,"" ],
            ["f","biaserrorX"           ,"" ],
            ["f","biaserrorY"           ,"" ],
            ["f","debugfloat"           ,"V" ],
            ["L","tmpdebugU32"          ,"" ]
        ]
        elif dataformat == 6:
            datastructure = [
            ["L","statusbits"           ,"" ],
            ["f","TECtemp"              ,"degC" ],
            ["f","MPD1"                 ,"mW"],
            ["L","uint8s"               ,""],
            ["f","biasvoltageXI"        ,"V" ],
            ["f","biasvoltageXQ"        ,"V" ],
            ["f","biasvoltageYI"        ,"V" ],
            ["f","biasvoltageYQ"        ,"V" ],
            ["f","biasvoltageX"         ,"V" ],
            ["f","biasvoltageY"         ,"V" ],
            ["f","biaserrorXI"          ,"deg" ],
            ["f","biaserrorXQ"          ,"deg" ],
            ["f","biaserrorYI"          ,"deg" ],
            ["f","biaserrorYQ"          ,"deg" ],
            ["f","biaserrorX"           ,"deg" ],
            ["f","biaserrorY"           ,"deg" ],
            ["f","lockscore"            ,"" ],
            ["f","debugfloat"           ,"" ],
            ["L","tmpdebugU32"          ,"" ]
        ]
        elif dataformat == 7:
            datastructure = [
            ["L","statusbits"           ,"" ],
            ["f","TECtemp"              ,"degC" ],
            ["f","MPD1"                 ,"mW"],
            ["L","uint8s"               ,""],
            ["f","biasvoltageXI"        ,"V" ],
            ["f","biasvoltageXQ"        ,"V" ],
            ["f","biasvoltageYI"        ,"V" ],
            ["f","biasvoltageYQ"        ,"V" ],
            ["f","biasvoltageX"         ,"V" ],
            ["f","biasvoltageY"         ,"V" ],
            ["f","biaserrorXI_I"        ,"deg" ],
            ["f","biaserrorXQ_I"        ,"deg" ],
            ["f","biaserrorYI_I"        ,"deg" ],
            ["f","biaserrorYQ_I"        ,"deg" ],
            ["f","biaserrorX_I"         ,"deg" ],
            ["f","biaserrorY_I"         ,"deg" ],
            ["f","lockscore"            ,"" ],
            ["f","debugfloat"           ,"" ],
            ["L","tmpdebugU32"          ,"" ]
        ]
        else:
            print "x99 Error: unrecognized data format"
            datastructure = defaultdatastructure  # returns None for these variables
            returnblank = True

        pendingbits = [
            [0, "biasscan"],
            [1, "stepping"],
            [2, "looprunning"]
        ]

        if len(returndat) is not len(datastructure) * 4:
            print 'incorrect return data length'

        masks = [
            [0, 'dithermask'],
            [1, 'adcmask'],
            [2, 'imagmask']
        ]

        class x99:
            """ """

            def __repr__(self):
                returnval = []
                padding = len(
                    max([a for a in dir(self) if not a.startswith('_')] + map(operator.itemgetter(1), datastructure), key=len))
                for vartype, varname, units in datastructure:
                    if varname not in ['pendingbits', 'statusbits', 'debugfloat', 'tmpdebugU32', 'uint8s', 'dataformat']:
                        returnval.append(
                            ' '.join([varname.rjust(padding), '=', str(round(getattr(self, varname), 4)), units]))
                if hasattr(x99dat, 'pendingbits'):
                    for pendingbit, varname in pendingbits:
                        returnval.append(' '.join([varname.rjust(padding), '=', str(getattr(self, varname))]))
                # namedvars=[]
                for varname in [a for a in dir(self) if not a.startswith('_') and
                                not a in map(operator.itemgetter(1), datastructure) and
                                not a in map(operator.itemgetter(1), pendingbits) and
                                not a in ['pendingbits', 'statusbits', 'debugfloat', 'tmpdebugU32']]:
                    returnval.append(' '.join([varname.rjust(padding), '=', str(getattr(self, varname))]))
                # returnval += sorted(namedvars)
                # returnval.append( 'tunerstate'.rjust(padding)+ ' = ' + self.tunerstate )
                retval2 = []
                for v, w in zip(returnval[0:int(math.ceil(len(returnval) / 2.0))],
                                returnval[int(math.ceil(len(returnval) / 2.0)):] + ['']):
                    retval2.append(' '.join([v.ljust(35), w]))
                return('\n'.join(['x99data format v' + str(self.dataformat) + ':'] + retval2))

        x99dat = x99()

        if not returnblank:
            pos = 0
            for vartype, varname, units in datastructure:
                rawdata = returndat[pos:pos + 4]
                # print "Decoding position:",pos," ".join("{:02X}".format(ord(c)) for c in rawdata)
                decodeddata = struct.unpack('<' + vartype, rawdata)[0]
                setattr(x99dat, varname, decodeddata)
                pos = pos + 4

            if hasattr(x99dat, 'pendingbits'):
                for pendingbit, varname in pendingbits:
                    setattr(x99dat, varname, (x99dat.pendingbits & (1 << pendingbit)) != 0)

            for maskbit, varname in masks:
                setattr(x99dat, varname, (ord(returndat[1]) & (1 << maskbit)) != 0)

            modA = ord(returndat[2])
            modB = ord(returndat[3])
            if modA < 7:
                setattr(x99dat, "modA", self.BIASNAMES[modA])
            if modB < 7:
                setattr(x99dat, "modB", self.BIASNAMES[modB])

            setattr(x99dat, "currentcontrolstep", ord(returndat[12]))
            setattr(x99dat, "alarms", ord(returndat[13]))
            setattr(x99dat, "mzstate", self.mzstate(ord(returndat[14])).name )
            setattr(x99dat, "dataformat", dataformat)
            
        else:
            if (self._dummyx99data is None):
                for vartype, varname, units in datastructure:
                    setattr(x99dat, varname, None)
                for pendingbit, varname in pendingbits:
                    setattr(x99dat, varname, None)
                self._dummyx99data = x99dat

        return(x99dat)

    def x99dict(self):
        """ """
        x99tmp = self.readx99()
        # remove unneeded entries
        for attribute in ['biaserrorXQ_I', 'biaserrorXI_I', 'biaserrorX_I',
                          'biaserrorYQ_I', 'biaserrorYI_I', 'biaserrorY_I',
                          'statusbits', 'pendingbits', 'uint8s']:
            if hasattr(x99tmp, attribute):
                delattr(x99tmp, attribute)

        return x99tmp.__dict__

    def getvoltagelist(self, x99data):
        """

        :param x99data: 

        """
        retdat = [None] * 6
        retdat[self.BIAS_XI] = x99data.biasvoltageXI
        retdat[self.BIAS_XQ] = x99data.biasvoltageXQ
        retdat[self.BIAS_X] = x99data.biasvoltageX
        retdat[self.BIAS_YI] = x99data.biasvoltageYI
        retdat[self.BIAS_YQ] = x99data.biasvoltageYQ
        retdat[self.BIAS_Y] = x99data.biasvoltageY
        return np.array(retdat)

    def setvoltagelist(self, voltagelist):
        """

        :param voltagelist: 

        """
        for bias in [self.BIAS_XI, self.BIAS_XQ, self.BIAS_X, self.BIAS_YI, self.BIAS_YQ, self.BIAS_Y]:
            self.setbiasvoltage(bias, voltagelist[bias])

    def calcminimodvoltages(self, x99dat):
        """

        :param x99dat: 

        """
        Vcom = 0.0
        setattr(x99dat, "biasvoltageIX1", Vcom + x99dat.biasvoltageXI + x99dat.biasvoltageX)
        setattr(x99dat, "biasvoltageIX2", Vcom - x99dat.biasvoltageXI + x99dat.biasvoltageX)
        setattr(x99dat, "biasvoltageQX1", Vcom + x99dat.biasvoltageXQ - x99dat.biasvoltageX)
        setattr(x99dat, "biasvoltageQX2", Vcom - x99dat.biasvoltageXQ - x99dat.biasvoltageX)
        setattr(x99dat, "biasvoltageIY1", Vcom + x99dat.biasvoltageYI + x99dat.biasvoltageY)
        setattr(x99dat, "biasvoltageIY2", Vcom - x99dat.biasvoltageYI + x99dat.biasvoltageY)
        setattr(x99dat, "biasvoltageQY1", Vcom + x99dat.biasvoltageYQ - x99dat.biasvoltageY)
        setattr(x99dat, "biasvoltageQY2", Vcom - x99dat.biasvoltageYQ - x99dat.biasvoltageY)
        return x99dat

    def biasscan(self, biasscanchannel, scantype,
                 startvoltage, endvoltage, numsteps, stepskips, toplot=False, log99=False, 
                 plotdemods=False, redline=True, quiet=False, imag=False):
        """Scans bias between specified voltage.  Returns voltages for local maxima and minima of specified signal (Amplitude, demodulation etc.

        :param biasscanchannel: 
        :param scantype: 
        :param startvoltage: 
        :param endvoltage: 
        :param numsteps: 
        :param stepskips: 
        :param toplot:  (Default value = False)
        :param log99:  (Default value = False)
        :param plotdemods:  (Default value = False)

        """

        # self.newevent(self.task.MZ_TASK, self.mzevent.MZ_STOP)
        # self.waitforstate( self.task.MZ_TASK, self.mzstate.MZ_LOWPOWER, timeout = 5, log99 = log99)
        initvoltages=self.getvoltagelist(self.readx99())

        self.writeuint32(0x17, biasscanchannel, index=0),
        self.writeuint32(0x17, scantype, index=1),  # scan type
        self.writefloat(0x17, startvoltage, index=2),  # start voltage
        self.writefloat(0x17, endvoltage, index=3),  # end voltage
        # self.writefloat( 0x17, 3, 0.0) # start voltage
        # self.writefloat( 0x17, 4, 2.49) # end voltage
        self.writeuint32(0x17, numsteps, index=4),  # steps
        self.writeuint32(0x17, stepskips, index=5),  # steps

        scanstart = time.time()
        self.newevent(self.task.MZ_TASK, self.mzevent.MZ_STARTBIASSCAN)
        if not self.waitforstate(self.task.MZ_TASK, self.mzstate.MZ_BIASSCAN, timeout=1.0, log99=log99):    
            if not self.checkinstate(self.task.MZ_TASK, self.mzstate.MZ_BIASSCAN):
                print "Unable to start bias scan"
                print self.getstates()
                return [None, None]
            # print self.writechar(0x17, 6, 0)  # go

        # pending = self.decodeuint32array(self.writereg(0x00))[0]

        snapshot = self.readx99()
        snapshot.time = time.time() - scanstart
        df = pd.DataFrame()
        df = df.append(snapshot.__dict__,ignore_index=True)

        count = 0
        sentvline = False
        
        while True:  # 1 != 0 or count < 5: # bit 1 is the bias scan pending mask
            count = count + 1
            # pending = self.decodeuint32array(self.writereg(0x00))[0]
            if log99:
                #x99info=self.readx99()
                snapshot = self.readx99()
                snapshot.time = time.time() - scanstart
                if imag:
                    snapshotimag = self.readx99(x99version=7)
                    snapshot.__dict__ = dict( snapshot.__dict__.items() + snapshotimag.__dict__.items() ) # TODO: mean of numericals
                df = df.append(snapshot.__dict__,ignore_index=True)
                
            voltages=self.getvoltagelist(snapshot)
            if redline and not sentvline and voltages[biasscanchannel] > initvoltages[biasscanchannel]:
                    self.logentry('Datalog redline ' + str(biasscanchannel))
                    sentvline = True
            # rawdata = self.decodeuint16array(writechar( 0x20, 5, 0))
            # voltage = rawdata[0] * ( 2.5 / 65535.0)
            # flags = rawdata[1]
            # biasvoltage = self.decodefloatarray(
            #    self.writereg(0x18))[biasscanchannel]
            if not quiet and not count % 100:
                sys.stdout.write('.')
            if not self.checkinstate(self.task.MZ_TASK, self.mzstate.MZ_BIASSCAN):
                break
            if plotdemods and not count % 400:
                self.plotdemods()
            # print pending, biasvoltage  # , flags
        if not quiet:
            sys.stdout.write('\n')

        scanend = time.time()
        if not quiet:
            print 'scanned in', scanend - scanstart, 'seconds'

        cpuusage = ord(self.writereg(0x11))

        # plt.plot(biasvoltages, linestyle='-', color='magenta')

        pdsvoltages = None
        biasvoltages = None

        if (toplot):
            scandat = self.writereg(0x21)

            pdsvoltages = self.decodeuint16array(scandat)

            biasvoltages = np.linspace(
                startvoltage,
                endvoltage,
                numsteps)

            print len(pdsvoltages), len(biasvoltages)

            plt.scatter(
                biasvoltages,
                pdsvoltages,
                linestyle='-',
                color='magenta',
                s=1)
            plt.show()

        if not quiet:
            print "CPU usage:", cpuusage, '%'
        return df
 
    def getbiasscanresults(self):
        """ """
        minmaxpoints = self.decodefloatarray(self.writereg(0x25))
        return pd.DataFrame(np.reshape(minmaxpoints, (len(minmaxpoints)/4,4)), columns=['minpos', 'min', 'maxpos','max']).dropna()

    def setbiasvoltage(self, bias, voltage):
        """Set specified bias channel voltage (power)

        :param bias: 
        :param voltage: 

        """
        return self.writefloat(0x14, voltage, index = bias)

    def getbiasvoltage(self, biaschannel=None):
        """

        :param biaschannel:  (Default value = None)

        """
        if biaschannel is None:
            print "Must specify bias channel"
            return
        return self.readfloat(regnum=0x13, index=biaschannel )
        
    def getbiasvoltages(self):
        """ """
        minmaxpoints = self.decodefloatarray(self.writereg(0x18))
        return pd.DataFrame(np.reshape(minmaxpoints, (len(minmaxpoints)/4,4)), columns=['minpos', 'min', 'maxpos','max']).dropna()

    def getmpd1pwr(self):
        """
        """
        return self.readfloat(regnum=0x20)

    def getmodulepwr(self):
        """
        """
        if  'self.ctrl.ADC.P3V3_ADC' in locals():
            return self._floatwithunits(self.ctrl.ADC.P3V3_ADC() * self.ctrl.ADC.MOD_I_ADC(), units='W')
        else:
            return self._floatwithunits(self.ctrl.ADC.ABC_AVDD_ADC() * self.ctrl.ADC.MOD_I_ADC(), units='W')
    
    def setmindemodaverages(self, mindemodaverages):
        """

        :param mindemodaverages: 

        """
        return self.writeuint16(0x33, mindemodaverages, index=1)



    def logtemperature(self):
        """ """
        if os.path.isfile(self._temperaturefilename):
            with open(self._temperaturefilename, 'r') as f:
                for line in f:
                    self.datalogentry(line)

    def readoma(self):
        """ """
        if os.path.isfile(self._omafilename):
            omadata = pd.read_table(self._omafilename, sep=' ', header=None, names=["Time", "variable", "value"])
            if self.lastomadata is not None:
                if self.lastomadata.equals(omadata):
                    return None  # don't log if not changed
            self.lastomadata = omadata
            return omadata
        else:
            return None

    def logoma(self, omadata=None):
        """

        :param omadata:  (Default value = None)

        """
        if omadata is None:
            omadata = self.readoma()
        for i, row in omadata.iterrows():
            self.datalogentry(str(row.Time) + ' ' + str(row.variable) + ' ' + str(row.value))
                              # todo: fix back-and-forth float to string conversions

    def readdemod(self, biaschannel=0xFF):
        """

        :param biaschannel:  (Default value = 0xFF)

        """
        rawdata = self.writechar(0x10, index=biaschannel, data=0)  # 0xFF for current channel
        uint32list = self.decodeuint32array(rawdata)
        floatlist = []
        countlist = []
        for mycount, mysum in [uint32list[i:i + 2] for i in range(0, len(uint32list), 2)]:
            countlist.append(mycount)
            if mycount == 0:
                floatlist.append(np.nan)
            else:
                floatlist.append(float(mysum) / (65535.0 * float(mycount)))
        return [countlist, floatlist]

    def readdedmoddf(self, maxloopstep=2):
        demoddat = [None] * maxloopstep
        for loopstep in range(maxloopstep):
            demoddat[loopstep] = pd.DataFrame(self.readdemod(loopstep)).transpose()[1]
            demoddat[loopstep].rename('loopstep'+str(loopstep),inplace=True)
        return pd.DataFrame(demoddat).transpose()

    def readmodrefs(self):
        """ """
        moddata = []
        for dacchannel in range(0, 6):
            rawdata = self.writereg(0x41, index=dacchannel)
            uint16list = self.decodeuint16array(rawdata)
            moddata.append([((x - 2048) / 2048.0) for x in uint16list])
        retdat = pd.DataFrame(map(list, zip(*moddata)), columns=self.BIASNAMES[:-1])
        retdat['index'] = retdat.index
        return retdat

    def getmodfreqs(self):
        """ """
        rawdata = self.writereg(0x44)
        floatlist = self.decodefloatarray(rawdata)

        return floatlist

    def plotmodrefs(self):
        """Plot the modulation reference waveforms"""
        modrefs = self.readmodrefs()
        melted = pd.melt(modrefs, id_vars=['index'])
        g = sns.FacetGrid(melted, row='variable', size=1.5, aspect=5)
        g = g.map(sns.pointplot, 'index', 'value', scale=0.5)
        # modplot = plt.plot(modrefs,marker='o', linestyle='-', markersize=4)
        # plt.legend(iter(modplot), self.BIASNAMES, loc='center left', bbox_to_anchor=(1, 0.5))
        # plt.show()
        return g

    def setmodshapealpha(self, alpha):
        """

        :param alpha: 

        """
        self.writefloat(0x27, alpha, index=1 )
        self.plotdemodrefs()

    def readdemodref(self, demodtype):
        """

        :param demodtype: 

        """
        rawdata = self.writechar(0x37, index=demodtype, data=0)
        # floatlist = self.decodeuSfpxparray(rawdata)
        floatlist = self.decodefloatarray(rawdata)
        #demodlength = len(floatlist)
        return floatlist

    def plotdemodref(self, biaschannel=0xFF):
        """Plot the demodulation reference waveforms

        :param biaschannel:  (Default value = 0xFF)

        """
        floatlist = self.readdemodref(biaschannel)
        plt.figure(figsize=(8,3))
        plt.plot(floatlist, marker='o', linestyle='-')
            # plt.ylim(0,4095)
        plt.show()
        print 'Mean:', np.nanmean(floatlist)

    def plotdemod(self, biaschannel=0xFF):
        """Plots the averaged PD waveforms (ADC data) for one bias
        
        :param biaschannel:  (Default value = 0xFF)

        """
        countlist, floatlist = self.readdemod(biaschannel)
        plt.plot(floatlist, marker='o', linestyle='-')
            # plt.ylim(0,4095)
        plt.show()
        print 'Mean:', np.nanmean(floatlist)

    def plotdemods(self, maxloopsteps=6):
        """Plots the averaged PD waveforms (ADC data) for each bias

        :param maxloopsteps:  (Default value = 6)

        """
        floatlists = []
        countlists = []
        for loopstep in range(0, maxloopsteps):
            # rawdata=self.writechar(0x10, biaschannel, 0)
            # floatlists.append([x / 65535.0 for x in self.decodefloatarray(rawdata)])
            countlist, floatlist = self.readdemod(loopstep)
            self.readx99()  # keep up to date.
            if np.sum(countlist) < 2:
                continue
            floatlists.append(floatlist)
            countlists.append(countlist)
        if not len(floatlists) > 1:
            print "No data returned for demods"
            return
        maxloopsteps = len(floatlists)
        x = range(0, len(floatlists[0]))
        miny = np.nanmin(floatlists)
        maxy = np.nanmax(floatlists)
        meany = np.nanmean(floatlists)
        ax = []
        f, ax = plt.subplots(maxloopsteps, figsize=(8, maxloopsteps * 1.2), sharex=True, sharey=False)
        for loopstep in range(0, maxloopsteps):
            ax[loopstep].plot(x, floatlists[loopstep], marker='o', linestyle='-', markersize=3)
            ax[loopstep].set_ylabel('Step ' + str(loopstep))
            ax[loopstep].set_ylim([miny, maxy])
            ax[loopstep].axhline(linewidth=0.5, y=meany, linestyle='--', color='black')
            ax[loopstep].get_xaxis().get_major_formatter().set_useOffset(False)
            ax[loopstep].yaxis.set_major_formatter(matplotlib.ticker.ScalarFormatter(useOffset=False))
        ax[0].set_title('TIA AC signal')
        ax[maxloopsteps - 1].set_xlim([-1, len(floatlists[0])])
        # Fine-tune figure; make subplots close to each other and hide x ticks for
        # all but bottom plot.
        f.subplots_adjust(hspace=0.15)
        plt.setp([a.get_xticklabels() for a in f.axes[:-1]], visible=False)
        # plt.plot(floatlists,marker='o', linestyle='-')
            # plt.ylim(0,4095)
        plt.show()
        print 'Counts:  Min: ', np.nanmin(countlists), "Max: ", np.nanmax(countlists)  # "Mean: ", np.nanmean(countlists),
        print 'Mean:', np.nanmean(floatlists)
        return self.adcstats()

    def plotdemodrefs(self):
        """ """
        floatlists = []
        #countlists = []
        #for demod in range(0, 6):
        for demod in [self.DEMOD_A, self.DEMOD_B, self.DEMOD_AB]:
            # rawdata=self.writechar(0x10, biaschannel, 0)
            # floatlists.append([x / 65535.0 for x in self.decodefloatarray(rawdata)])
            floatlist = self.readdemodref(demod)
            floatlists.append(floatlist)
        x = range(0, len(floatlists[0]))
        #f, (ax1, ax2, ax3, ax4, ax5, ax6) = plt.subplots(6, figsize=(8, 8), sharex=True, sharey=False)
        f, (ax1, ax2, ax4 ) = plt.subplots(3, figsize=(8, 4), sharex=True, sharey=False)
        ax1.plot(x, floatlists[0], marker='o', linestyle='-', markersize=3)
        ax2.plot(x, floatlists[1], marker='o', linestyle='-', markersize=3)
        #ax3.plot(x, floatlists[2], marker='o', linestyle='-', markersize=3)
        ax4.plot(x, floatlists[2], marker='o', linestyle='-', markersize=3)
        #ax5.plot(x, floatlists[4], marker='o', linestyle='-', markersize=3)
        #ax6.plot(x, floatlists[5], marker='o', linestyle='-', markersize=3)
        ax1.set_ylabel('DEMOD_A')
        ax2.set_ylabel('DEMOD_B')
        #ax3.set_ylabel('DEMOD_C')
        ax4.set_ylabel('DEMOD_AB')
        #ax5.set_ylabel('DEMOD_AC')
        #ax6.set_ylabel('DEMOD_BC')
        ax1.set_title('Demod References')
        ax4.set_xlim([-1, len(floatlists[0])])
        miny = np.nanmin(floatlists)
        maxy = np.nanmax(floatlists)
        meany = np.nanmean(floatlists)
        ax1.set_ylim([miny, maxy])
        ax2.set_ylim([miny, maxy])
        #ax3.set_ylim([miny, maxy])
        ax4.set_ylim([miny, maxy])
        #ax5.set_ylim([miny, maxy])
        #ax6.set_ylim([miny, maxy])
        ax1.axhline(linewidth=0.5, y=meany, linestyle='--', color='black')
        ax2.axhline(linewidth=0.5, y=meany, linestyle='--', color='black')
        #ax3.axhline(linewidth=0.5, y=meany, linestyle='--', color='black')
        ax4.axhline(linewidth=0.5, y=meany, linestyle='--', color='black')
        #ax5.axhline(linewidth=0.5, y=meany, linestyle='--', color='black')
        #ax6.axhline(linewidth=0.5, y=meany, linestyle='--', color='black')
        # Fine-tune figure; make subplots close to each other and hide x ticks for
        # all but bottom plot.
        f.subplots_adjust(hspace=0.15)
        plt.setp([a.get_xticklabels() for a in f.axes[:-1]], visible=False)
        # plt.plot(floatlists,marker='o', linestyle='-')
            # plt.ylim(0,4095)
        plt.show()
        print 'Mean:', np.nanmean(floatlists)

    def adcstats(self, quiet=False):
        """ """
        adcstats = self.decodeuint16array(self.writechar(0x30, index=0, data=0))
        if not quiet:
            print 'ADC stats (since last read)', adcstats, np.mean(adcstats)
            print 'Saturation', np.diff(adcstats)[0] / 65535.0
            if np.diff(adcstats)[0] / 65535.0 < -0.01:
                print "Stats incorrect, possibly ADC not running"
        return adcstats

    def plotadclogdata(self, numavr=1):
        """

        :param numavr:  (Default value = 1)

        """
        floatlists = []
        for i in range(numavr):
            rawdata = self.writechar(0x28, index=3, data=1)
            if (len(rawdata) == 2 and rawdata == 'NO'):
                print 'Data not available.  Please check FW compiled with ADCLOGGINGENABLED'
                return
            floatlists.append([x / 65535.0 for x in self.decodeuint16array(rawdata)])
            sys.stdout.write('.')
            sys.stdout.flush()
            time.sleep(np.random.uniform(0.0, 0.05))
        # plt.figure(num=None, figsize=(10, 6))
        # plt.plot(floatlist, linestyle='-', linewidth=0.1)
            # plt.ylim(0,4095)
        # plt.show()
        # plt.figure(num=None, figsize=(10, 6))
        # plt.semilogy(np.abs(np.fft.fft(floatlist)), linestyle='-', linewidth=0.1)
            # plt.ylim(0,4095)
        # plt.show()
        ffts = []
        for i in range(numavr):
            y = floatlists[i]

            Fs = 80000.0
            # sampling rate
            Ts = 1.0 / Fs
            # sampling interval

            n = len(y)  # length of the signal
            t = np.arange(n) * Ts  # time vector
            k = np.arange(n)
            T = n / Fs
            frq = k / T  # two sides frequency range
            frq = frq[range(n / 2)]  # one side frequency range

            Y = np.fft.fft(y) / n  # fft computing and normalization
            Y = Y[range(n / 2)]
            ffts.append(Y)

        fftavr = np.mean(np.abs(ffts), axis=0)

        fig, ax = plt.subplots(4, 1, figsize=(12, 12))
        # ax[0].plot(t,y,linewidth=0.5,marker='o',markersize=3)
        ax[0].plot(t, y, linewidth=0.1)
        ax[0].set_xlabel('Time')
        ax[0].set_ylabel('Amplitude')
# ax[1].semilogy(frq/1000.0,abs(Y),'r',linewidth=0.2) # plotting the spectrum
# ax[1].plot(frq/1000.0,abs(Y),'r',linewidth=0.2) # plotting the spectrum
# ax[1].set_ylim([0,0.02])
#        ax[1].set_xlabel('Freq (kHz)')
#        ax[1].set_ylabel('|Y(freq)|')
        ax[1].semilogy(frq[1:] / 1000.0, abs(fftavr[1:]), 'r', linewidth=0.2)  # plotting the spectrum
        # ax[1].plot(frq/1000.0,abs(Y),'r',linewidth=0.2) # plotting the spectrum
        # ax[1].set_ylim([0,0.02])
        ax[1].set_xlabel('Freq (kHz)')
        ax[1].set_ylabel('avr |Y(freq)|')
        ax[2].semilogy(frq[1:len(frq) / 5], fftavr[1:len(fftavr) / 5], 'r', linewidth=0.5)  # plotting the spectrum
        ax[2].set_xlabel('Freq (Hz)')
        ax[2].set_ylabel('avr |Y(freq)|')
        ax[3].semilogy(frq[1:len(frq) / 50], fftavr[1:len(fftavr) / 50], 'r', linewidth=0.5)  # plotting the spectrum
        ax[3].set_xlabel('Freq (Hz)')
        ax[3].set_ylabel('avr |Y(freq)|')
        # ax[3].axvline(x=179.0)
        fig.show()
        print 'Note: DC removed'
        return(floatlists, [fftavr])

    def setmask(self, index, value):
        """

        :param index: 
        :param value: 

        """
        self.writechar(0x28, index, value)

    def find_nearest(self, array, value):
        """

        :param array: 
        :param value: 

        """
        array = [x for x in array if str(x[0]) != 'nan']
        idx = np.array([abs(i[0] - value) for i in array]).argmin()
        return array[idx]

    def find_nearest_ndx(self, array, value):
        """

        :param array: 
        :param value: 

        """
        array = [x for x in array if str(x[0]) != 'nan']
        idx = np.array([abs(i[0] - value) for i in array]).argmin()
        return idx

    def setminmaxquad(self, bias, targettype, startvoltage=10, endvoltage=40, numsteps=500, stepskips=1, quiet=False):
        """

        :param bias: 
        :param targettype: 
        :param startvoltage:  (Default value = -10)
        :param endvoltage:  (Default value = 10)
        :param numsteps:  (Default value = 500)
        :param stepskips:  (Default value = 1)

        """
        #self.vlinecounter = self.vlinecounter + 1
        #self.logentry('Datalog vline ' + str(self.vlinecounter))
        self.biasscan(biasscanchannel=bias,
                      scantype=self.DCPHOTODIODE,
                      startvoltage=startvoltage,
                      endvoltage=endvoltage,
                      numsteps=numsteps,
                      stepskips=stepskips,
                      toplot=False,  # only use if STORESCAN defined in fw
                      log99=True,
                      redline=False,
                      quiet=quiet)

        results = self.getbiasscanresults()
        if not quiet:
            print results

        minmaxpoints = self.decodefloatarray(self.writereg(0x25))

        mins = [minmaxpoints[idx:idx + 2] for idx in range(0, len(minmaxpoints), 4)]
        maxs = [minmaxpoints[idx + 2:idx + 4] for idx in range(0, len(minmaxpoints), 4)]

        #print 'Minima:'
        #pprint.pprint(mins)
        #print 'Maxima:'
        #pprint.pprint(maxs)

        target =  (startvoltage + endvoltage / 2.0)


        maxpoint = results.ix[ results['max'].idxmax(), 'maxpos']
        minpoint = results.ix[ results['min'].idxmax(), 'minpos']

        #maxpoint = maxs[self.find_nearest_ndx(maxs, target)][0]
        #minpoint = mins[self.find_nearest_ndx(mins, target)][0]
        vPi = abs(maxpoint - minpoint)
        midpoint = (maxpoint + minpoint) / 2.0
        midpointpos = maxpoint > minpoint

        # self.monitor(2,False,False)
        if not quiet:
            print 'V_Pi is', vPi, '(', maxs[self.find_nearest_ndx(maxs, target)][0], '-', mins[self.find_nearest_ndx(mins, 0.0)][0], ')'

        if targettype == self.MINPOINT:
            if not quiet:
                print 'Setting', self.BIASNAMES[bias], 'to min at', minpoint, 
            self.setbiasvoltage(bias, minpoint)
        elif targettype == self.MAXPOINT:
            if not quiet:
                print 'Setting', self.BIASNAMES[bias], 'to max at', maxpoint, 
            self.setbiasvoltage(bias, maxpoint)
        elif targettype == self.QUADPOINTPOS:
            if midpointpos:
                if not quiet:
                    print 'Setting', self.BIASNAMES[bias], 'to quad+ at', midpoint, 
                self.setbiasvoltage(bias, midpoint)
            else:
                if (midpoint > target):  # shift, but keep close to middle
                    midpoint -= vPi
                else:
                    midpoint += vPi
                if not quiet:
                    print 'Setting', self.BIASNAMES[bias], 'to quad+ at', midpoint
                self.setbiasvoltage(bias, midpoint)
        elif targettype == self.QUADPOINTNEG:
            if not midpointpos:
                if not quiet:
                    print 'Setting', self.BIASNAMES[bias], 'to quad- at', midpoint
                self.setbiasvoltage(bias, midpoint)
            else:
                if (midpoint > target):  # shift, but keep close to middle
                    midpoint -= vPi
                else:
                    midpoint += vPi
                if not quiet:
                    print 'Setting', self.BIASNAMES[bias], 'to quad- at', midpoint
                self.setbiasvoltage(bias, midpoint)

        # self.monitor(1,False,False)

        return [vPi, mins, maxs, midpoint, midpointpos]

    def getuptime(self, remote=False):
        return self.decodeuint32array(self.readreg(0x43,remote=remote))[0] / 1000.0

    def monitor(self, monitortime, readoma=False, plotdemods=False, maxloopsteps=6, settomeans=False, quiet=False, nthdata=1):
        """

        :param monitortime: 
        :param readoma:   (Default value = False)
        :param plotdemods:   (Default value = False)
        :param maxloopsteps:  (Default value = 6)
        :param settomeans:  (Default value = False)
        :param quiet:  (Default value = False)

        """
        # averages bias voltages over monitortime
        j = 0
        i = 0
        n = 0
        omacount = 0
        starttime = time.time()
        omareadtime = time.time()
        biasvsums = [0.0] * 6
        biaserrsums = [0.0] * 6
        omadataall = pd.DataFrame()
        while time.time() - starttime < monitortime:
            i = i + 1
            if not i % 100 and not quiet:
                sys.stdout.write('.')
                sys.stdout.flush()
                #print self.getuptime(remote=True), self.getuptime(remote=False)
            if not i % nthdata:
                x99info = self.readx99()
                biasvsums[self.BIAS_XI] += x99info.biasvoltageXI
                biasvsums[self.BIAS_XQ] += x99info.biasvoltageXQ
                biasvsums[self.BIAS_YI] += x99info.biasvoltageYI
                biasvsums[self.BIAS_YQ] += x99info.biasvoltageYQ
                biasvsums[self.BIAS_X] += x99info.biasvoltageX
                biasvsums[self.BIAS_Y] += x99info.biasvoltageY
                biaserrsums[self.BIAS_XI] += x99info.biaserrorXI
                biaserrsums[self.BIAS_XQ] += x99info.biaserrorXQ
                biaserrsums[self.BIAS_YI] += x99info.biaserrorYI
                biaserrsums[self.BIAS_YQ] += x99info.biaserrorYQ
                biaserrsums[self.BIAS_X] += x99info.biaserrorX
                biaserrsums[self.BIAS_Y] += x99info.biaserrorY
                n += 1
            # if not i % 10:
            if time.time() - omareadtime > 1:
                omareadtime = time.time()
                if readoma:
                    omadata = self.readoma()
                    if omadata is not None:
                        omacount += 1
                        self.logoma(omadata)
                        omadata.Time = np.mean(omadata.Time)
                        omadataall = omadataall.append(omadata.pivot(index='Time', columns='variable', values='value'))

                # mz.logentry('OMA: '+str(oma.getEVM(0)) + ' ' + str(oma.getIQoffset(0)) + ' ' + str(oma.getQuadErr(0)) + ' ' +
                #            str(oma.getEVM(1)) + ' ' + str(oma.getIQoffset(1)) + ' ' + str(oma.getQuadErr(1)))
                j = j + 1
                if not j % 3 and not quiet:
                    print
                    print self.getstates()

                if not j % 6:
                    self.logtemperature()
                    if not quiet:
                        print
                        print (time.time() - starttime) / 60, 'mins'
                        # print self.calcminimodvoltages(x99info)
                        print x99info
                        print 'CPU usage', self.getcpuusage(), '%'
#                        print self.ctrl.ADC.ABC_AVDD_ADC()
#                        print self.ctrl.ADC.P3V3_ADC(), self.ctrl.ADC.P6V_ADC()


            if plotdemods and not i % 500:
                if not quiet:
                    self.plotdemods(maxloopsteps=maxloopsteps)

            omasummary = pd.DataFrame()

            if omacount > 0:
                try:
                    omasummary = pd.DataFrame({
                                             'mean': omadataall.mean(),
                                             'meanabs': omadataall.abs().mean(),
                                             'std': omadataall.std(),
                                             'stdabs': omadataall.abs().std(),
                                             'quantile01': omadataall.quantile(0.01),
                                             'quantile99': omadataall.quantile(0.99),
                                             'quantile99abs': omadataall.abs().quantile(0.99)}).stack().reset_index()
                    omasummary = omasummary.rename(columns={0: 'value'})
                    omasummary.index = omasummary["variable"].map(str) + '_' + omasummary["level_1"]
                    omasummary = omasummary.drop(['variable', 'level_1'], 1)
                    omasummary = omasummary.append(pd.DataFrame([time.time() - starttime, omacount], 
                                                                 index=['duration', 'n'], columns=['value']))
                except:
                    omasummary = pd.DataFrame()

        meanbiases = [x / float(n) for x in biasvsums]
        if settomeans:
            # self.writefloat(0x27, 0, 0) # stepscale to zero, to stop adjustment before setting means
            self.newevent(self.task.MZ_TASK, self.mzevent.MZ_PAUSEADJUSTMENT)
            self.setbiasvoltage(self.BIAS_XI, meanbiases[self.BIAS_XI])
            self.setbiasvoltage(self.BIAS_XQ, meanbiases[self.BIAS_XQ])
            self.setbiasvoltage(self.BIAS_X, meanbiases[self.BIAS_X])
            self.setbiasvoltage(self.BIAS_YI, meanbiases[self.BIAS_YI])
            self.setbiasvoltage(self.BIAS_YQ, meanbiases[self.BIAS_YQ])
            self.setbiasvoltage(self.BIAS_Y, meanbiases[self.BIAS_Y])
        return ([meanbiases,
                 [x / float(n) for x in biaserrsums],
                 omasummary.transpose(), omadataall])  # return the means

    def dotest( self,
                testtime=60,
                loop=1,
                dithermag=0.5,
                stepscale=20.0,
                settomeans=False,
                validdata=0,
                readoma=False,
                plotdemods=True,
                idnum=None,
                maxloopsteps=6):
        """

        :param testtime:  (Default value = 60)
        :param loop:  (Default value = 1)
        :param dithermag:  (Default value = 0.5)
        :param stepscale:  (Default value = 20.0)
        :param settomeans:  (Default value = False)
        :param validdata:  (Default value = 0)
        :param readoma:  (Default value = False)
        :param plotdemods:  (Default value = True)
        :param idnum:  (Default value = None)
        :param maxloopsteps:  (Default value = 6)

        """
        self.logentry('Datalog stepscale NaN')  # NaNs break the line
        self.logentry('Datalog loop NaN')
        self.logentry('Datalog vline NaN')
        self.logentry('Datalog idnum NaN')
        self.modamplitude(self.MOD_A, dithermag)
        self.modamplitude(self.MOD_B, dithermag)
        self.setloopstepscale(stepscale)
        # self.setmask(self.MASKLOOPADJUSTMENT,False)
        if loop:
            print self.newevent(self.task.MZ_TASK, self.mzevent.MZ_STARTTESTLOOP)
        else:
            print self.newevent(self.task.MZ_TASK, self.mzevent.MZ_STOP)
        self.logentry('Datalog dithermag ' + str(dithermag))
        self.logentry('Datalog stepscale ' + str(stepscale))
        self.logentry('Datalog loop ' + str(loop))
        if idnum is not None:
            self.logentry('Datalog idnum ' + str(idnum))
        self.logentry('Datalog validdata 0')
        self.logentry('Datalog validdata ' + str(validdata))
        [meanbiases,
         meanerrors,
         omastats,
         omadataall] = self.monitor( testtime,
                                     readoma=readoma,
                                     plotdemods=plotdemods,
                                     maxloopsteps=maxloopsteps)
        self.logentry('Datalog validdata ' + str(validdata))
        self.logentry('Datalog validdata 0')
        if idnum is not None:
            self.logentry('Datalog idnum ' + str(idnum))
        self.logentry('Datalog dithermag ' + str(dithermag))
        self.logentry('Datalog stepscale ' + str(stepscale))
        self.logentry('Datalog loop ' + str(loop))
        self.logentry('Datalog vline ' + str(self.vlinenum))
        self.vlinenum += 1
        omastats['loop'] = loop
        omastats['dithermag'] = dithermag
        omastats['stepscale'] = stepscale
        if settomeans:
            # self.writefloat(0x27, 0, 0) # stepscale to zero, to stop adjustment before setting means
            self.newevent(self.task.MZ_TASK, self.mzevent.MZ_PAUSEADJUSTMENT)
            self.setbiasvoltage(self.BIAS_XI, meanbiases[self.BIAS_XI])
            self.setbiasvoltage(self.BIAS_XQ, meanbiases[self.BIAS_XQ])
            self.setbiasvoltage(self.BIAS_X, meanbiases[self.BIAS_X])
            self.setbiasvoltage(self.BIAS_YI, meanbiases[self.BIAS_YI])
            self.setbiasvoltage(self.BIAS_YQ, meanbiases[self.BIAS_YQ])
            self.setbiasvoltage(self.BIAS_Y, meanbiases[self.BIAS_Y])
        return [meanbiases, meanerrors, omastats, omadataall]

    def initcoarseloop(self):
        """ """
        return self.writeuint32( regnum = 0x35, index=13, index2=0, val=0)

    def initfineloop(self):
        """ """
        return self.writeuint32( regnum = 0x35, index=14, index2=0, val=0)

    def initlockedloop(self):
        """ """
        return self.writeuint32( regnum = 0x35, index=15, index2=0, val=0)

    def loopcontrol(self, command):
        """

        :param command: 

        """
        self.writeuint32(0x32,  command, index=0)

    def setloopstepscale(self, stepscale):
        """

        :param stepscale: 

        """
        self.writefloat(0x27,  stepscale, index=0)

    def modphase(self, ditherchannel= None, phase=None):
        """

        :param ditherchannel:  (Default value = None)
        :param phase:  (Default value = None)

        """
        if ditherchannel is None:
            print "Must specify dither channel"
            return
        if phase is not None:
            return self.writefloat(regnum=0x23, data=phase,  index=ditherchannel )
        else:
            phase = self.readfloat(regnum=0x23, index=ditherchannel )
            phase.units='Rad'        
            return phase

    def demodphase(self, ditherchannel= None, phase=None):
        """

        :param ditherchannel:  (Default value = None)
        :param phase:  (Default value = None)

        """
        if ditherchannel is None:
            print "Must specify dither channel"
            return
        if phase is not None:
            return self.writefloat(regnum=0x13, data=phase,  index=ditherchannel )
        else:
            return self.readfloat(regnum=0x13, index=ditherchannel, units='Rad' )

    def modamplitude(self, ditherchannel=None, ditheramp=None):
        """

        :param ditherchannel:  (Default value = None)
        :param ditheramp:  (Default value = None)

        """
        if ditherchannel is None:
            print "Must specify dither channel"
            return
        if ditheramp is not None:
            return self.writefloat(regnum=0x19, data=ditheramp,  index=ditherchannel )
        else:
            return self.readfloat(regnum=0x19, index=ditherchannel )

    def loopsetup(self, loopstep,
                  pol,
                  modAsubmz,
                  modBsubmz,
                  modCsubmz,
                  submzIdemodsource,
                  submzQdemodsource,
                  submzPdemodsource,
                  submzItargetdemod,
                  submzQtargetdemod,
                  submzPtargetdemod,
                  submzIscale,
                  submzQscale,
                  submzPscale):
        """

        :param loopstep: 
        :param pol: 
        :param modAsubmz: 
        :param modBsubmz: 
        :param modCsubmz: 
        :param submzIdemodsource: 
        :param submzQdemodsource: 
        :param submzPdemodsource: 
        :param submzItargetdemod: 
        :param submzQtargetdemod: 
        :param submzPtargetdemod: 
        :param submzIscale: 
        :param submzQscale: 
        :param submzPscale: 

        """
        self.setloop_pol(loopstep, pol)
        self.setloop_modAsubmz(loopstep, modAsubmz)
        self.setloop_modBsubmz(loopstep, modBsubmz)
        #self.setloop_modCsubmz(loopstep, modCsubmz) # not implemented
        self.setloop_submzIdemodsource(loopstep, submzIdemodsource)
        self.setloop_submzQdemodsource(loopstep, submzQdemodsource)
        self.setloop_submzPdemodsource(loopstep, submzPdemodsource)
        self.setloop_submzItargetdemod(loopstep, submzItargetdemod)
        self.setloop_submzQtargetdemod(loopstep, submzQtargetdemod)
        self.setloop_submzPtargetdemod(loopstep, submzPtargetdemod)
        self.setloop_submzIscale(loopstep, submzIscale)
        self.setloop_submzQscale(loopstep, submzQscale)
        self.setloop_submzPscale(loopstep, submzPscale)
        # note: the last loop step is idenfitied with a pol of POL_NONE

    def setloop_pol(self, loopstep, pol):
        """

        :param loopstep: 
        :param pol: 

        """
        return self.writeuint32(0x35, pol, index=0, index2=loopstep)

    def setloop_modAsubmz(self, loopstep, modAsubmz):
        """

        :param loopstep: 
        :param modAsubmz: 

        """
        return self.writeuint32(0x35, modAsubmz, index=1, index2=loopstep)

    def setloop_modBsubmz(self, loopstep, modBsubmz):
        """

        :param loopstep: 
        :param modBsubmz: 

        """
        return self.writeuint32(0x35, modBsubmz, index=2, index2=loopstep)

    def setloop_modCsubmz(self, loopstep, modCsubmz):
        """

        :param loopstep: 
        :param modCsubmz: 

        """
        return self.writeuint32(0x35, modCsubmz, index=3, index2=loopstep)

    def setloop_submzIdemodsource(self, loopstep, submzIdemodsource):
        """

        :param loopstep: 
        :param submzIdemodsource: 

        """
        return self.writeuint32(0x35, submzIdemodsource, index=4, index2=loopstep)

    def setloop_submzQdemodsource(self, loopstep, submzQdemodsource):
        """

        :param loopstep: 
        :param submzQdemodsource: 

        """
        return self.writeuint32(0x35, submzQdemodsource, index=5, index2=loopstep)

    def setloop_submzPdemodsource(self, loopstep, submzPdemodsource):
        """

        :param loopstep: 
        :param submzPdemodsource: 

        """
        return self.writeuint32(0x35, submzPdemodsource, index=6, index2=loopstep)

    def setloop_submzItargetdemod(self, loopstep, submzItargetdemod):
        """

        :param loopstep: 
        :param submzItargetdemod: 

        """
        return self.writefloat(0x35, submzItargetdemod, index=7, index2=loopstep)

    def setloop_submzQtargetdemod(self, loopstep, submzQtargetdemod):
        """

        :param loopstep: 
        :param submzQtargetdemod: 

        """
        return self.writefloat(0x35, submzQtargetdemod, index=8, index2=loopstep)

    def setloop_submzPtargetdemod(self, loopstep, submzPtargetdemod):
        """

        :param loopstep: 
        :param submzPtargetdemod: 

        """
        return self.writefloat(0x35, submzPtargetdemod, index=9, index2=loopstep)

    def setloop_submzIscale(self, loopstep, submzIscale):
        """

        :param loopstep: 
        :param submzIscale: 

        """
        return self.writefloat(0x35, submzIscale, index=0, index2=loopstep)

    def setloop_submzQscale(self, loopstep, submzQscale):
        """

        :param loopstep: 
        :param submzQscale: 

        """
        return self.writefloat(0x35, submzQscale, index=11, index2=loopstep)

    def setloop_submzPscale(self, loopstep, submzPscale):
        """

        :param loopstep: 
        :param submzPscale: 

        """
        return self.writefloat(0x35, submzPscale, index=12, index2=loopstep)

    def bias_p_gain(self, biascontrolndx, target=None):
        """

        :param biascontrolndx: 
        :param target: (Default target = None)

        """
        if target is not None:
            return self.writefloat(0x36, target, index=5, index2=biascontrolndx)
        else:
            return self.decodefloatarray(self.readreg(0x36,index=5, index2=biascontrolndx))[0]

    def bias_i_gain(self, biascontrolndx, target):
        """

        :param biascontrolndx: 
        :param target: 

        """
        return self.writefloat(0x36, target, index=0, index2=biascontrolndx)

    def bias_d_gain(self, biascontrolndx, target):
        """

        :param biascontrolndx: 
        :param target: 

        """
        return self.writefloat(0x36, target, index=1, index2=biascontrolndx)

    def bias_clearisum(self, biascontrolndx):
        """

        :param biascontrolndx: 

        """
        return self.writefloat(0x36, 0.0, index=2, index2=biascontrolndx)

    def bias_isumlimit(self, biascontrolndx, target):
        """

        :param biascontrolndx: 
        :param target: 

        """
        return self.writefloat(0x36, target, index=3, index2=biascontrolndx)

    def bias_EMAalpha(self, biascontrolndx, target):
        """

        :param biascontrolndx: 
        :param target: 

        """
        return self.writefloat(0x36, target, index=4, index2=biascontrolndx)

    def bias_pidtarget(self, biascontrolndx, target):
        """

        :param biascontrolndx: 
        :param target: 

        """
        return self.writefloat(0x36, target, index=6, index2=biascontrolndx)

    def scanpterm(self, biastoscan=None, scanmin=-8000.0, scanmax=8000.0):
        """

        :param biastoscan:  (Default value = None)
        :param scanmin:  (Default value = -8000.0)
        :param scanmax:  (Default value = 8000.0)

        """

        self.bias_pidtarget(self.BIAS_XI, 0.0)
        self.bias_pidtarget(self.BIAS_XQ, 0.0)
        self.bias_pidtarget(self.BIAS_X, 0.0)
        self.bias_pidtarget(self.BIAS_YI, 0.0)
        self.bias_pidtarget(self.BIAS_YQ, 0.0)
        self.bias_pidtarget(self.BIAS_Y, 0.0)

        # biastoscan=self.BIAS_YQ
        # scanmin = -8000.0
        # scanmax = 8000.0

        scanstep = (scanmax - scanmin) / 10000.0

        omareadtime = time.time()
        i = 0
        for offset in np.arange(scanmin, scanmax, scanstep):
            i = i + 1
            if not i % 100:
                sys.stdout.write('.')
                sys.stdout.flush()
            self.readx99()
            self.setloop_pidtarget(biastoscan, offset)
            self.logentry('Datalog logdebug ' + str(offset))
            if abs(offset) < scanstep / 2.0:
                self.logentry('Datalog vline ' + str(15))
            if time.time() - omareadtime > 1:
                omareadtime = time.time()
                omadata = self.readoma()
                if omadata is not None:
                    self.logoma(omadata)

        self.bias_pidtarget(self.BIAS_XI, 0.0)
        self.bias_pidtarget(self.BIAS_XQ, 0.0)
        self.bias_pidtarget(self.BIAS_X, 0.0)
        self.bias_pidtarget(self.BIAS_YI, 0.0)
        self.bias_pidtarget(self.BIAS_YQ, 0.0)
        self.bias_pidtarget(self.BIAS_Y, 0.0)
        

    def connect(self, port=3, baud=115200, time_out=2.0):
        """Connect to the unit via UART

        :param port:  (Default value = 3)
        :param baud:  (Default value = 115200)
        :param time_out:  (Default value = 0.5)

        """

            
        if isinstance(port, basestring):
            self._port = port
        elif isinstance(port, int):
            if os.name == 'posix':
                self._port = "/dev/ttyUSB" + str(port)
            else:
                self._port = "com" + str(port)

        if self._OIFlogging:
            # with open(self._commslogfilename, "a") as oiflogfile:
            oiflogfile = open(self._commslogfilename, "a")
            oiflogfile.write(
                "# connecting to port " + str(
                    port) + " at " + time.strftime(
                        "%a %d %b %Y %H:%M:%S %Z") + "\n")
            oiflogfile.flush()
        else:
            if self._print_once:
                self._print_once = False
                print self.hilite("Warning: Not logging OIF transactions. Use\n mz.logfile(<filename>) to set log filename and\n mz.logging(True) to enable logging", False, True)


        if hasattr(port,'port'):
            print type(self).__name__,'connecting using existing port handle',
            self._ttyhandle = port
            port = self._ttyhandle.port
            self.baud = self._ttyhandle.baudrate
            print 'on', port,'at',self.baud
        else:            
            self._baud = baud
    
    
            print "Connecting to", self._port, "at", self._baud, "baud"
            try:
                self._ttyhandle = serial.Serial(self._port, baud, timeout=time_out)
            except serial.SerialException as e:
                print "Error:", e
                return

        buildstring =  self.buildstring()
        print buildstring
        if '; OSFP ;' in buildstring:  # fixme: add report of HW type to register
            self._moduletype = 'OSFP_ACO'
        else:
            self._moduletype = 'CFP2ACO'
        print "module_type: " + self._moduletype


        print 'test1'
        
        self.setstatemachinemapfromunit()

        self.setctrlfunctionsfromunit()
                
        if not hasattr(self,'ctrl'):  # should never happen.  If fails, try the old way.
            self._setupctrlfunctions_defunct()
        self.setcalmapfromunit()
        
        self.load_kv_from_module()
        
        # open the serial port
        if self._ttyhandle.isOpen():
#            self.t.save_handle(self._ttyhandle)
#            self.it.save_handle(self._ttyhandle)
            self._connectime = time.time()
            # self.buildstring()
            #self.io.connect()
        else:
            print "Unable to connect"
            self._port = None


    def _tsave(self, destination = 'TTM', bin = None, autoresponse = False, \
             verify = False, verify_retry = 0, skipclosevalues=False, defaultstonone=False, rawcal=False):
        
        if ('TTM'.startswith(destination.upper())):
            if (bin != None):
                raise 'Error: Bin must be specified as None when saving to the TTM'
            print 'Saving kv file to module',
            sys.stdout.flush() 
            time.sleep(0.001)       # allow time for printing              
            self.save_kv_to_module(writetoflash = True, skipclosevalues=skipclosevalues, rawcal=rawcal)
            return

        if ('FILE'.startswith(destination.upper())):

            if (type(bin) != type('hello')):

                raise 'Error: Bin must be specified as a string.'

            sys.stdout.flush()
            return(self.save_kv_to_file(filename=bin, defaultstonone=defaultstonone, rawcal=rawcal))

        raise 'Error: \'%s\' is not a valid source.' % (destination)
        

    def _trestore(self, source = 'TTM', bin = None, autoresponse = False, loadall=False, blank=False, defaultstonone=False, rawcal=False):

        if ('TTM'.startswith(source.upper())):

            if (bin != None):
                raise 'Error: Bin must be specified as None.'

            self._get_cal_map_ifneeded()

            if blank:
                sys.stderr.flush()
                sys.stderr.write('Creating blank kv file')
                sys.stderr.flush()
                time.sleep(0.001)       # allow time for printing
                if rawcal:
                    self.loadblankkv(calkvdict=self.cal.rawcaldict, kvdict=self.rawcal.kvdict)
                    self.fixkvorder(self.rawcal,'rawcaldict')
                else:
                    self.loadblankkv()
                    self.fixkvorder(self.kv)
                return

            if rawcal:
                self.load_kv_from_module(calcontainer=self.cal, calbase='rawcaldict', defaultstonone=defaultstonone)
                return

            sys.stderr.flush()
            sys.stderr.write('Retrieving kv file from board')
            sys.stderr.flush()
            time.sleep(0.001)       # allow time for printing              
            self.load_kv_from_module(calcontainer=self.cal, defaultstonone=defaultstonone)
            if loadall:
                sys.stderr.flush()
                sys.stderr.write('Retrieving kv defaults from nITLA')
                sys.stderr.flush()
                time.sleep(0.001)       # allow time for printing              
                self.load_kv_from_module(calcontainer=self.cal_defaults, defaultstonone=defaultstonone)
                sys.stderr.flush()
                sys.stderr.write('Retrieving kv flash values from nITLA')
                sys.stderr.flush()
                time.sleep(0.001)       # allow time for printing              
                self.load_kv_from_module(calcontainer=self.cal_flash, defaultstonone=defaultstonone)
            return

        if ('FILE'.startswith(source.upper())):

            self._get_cal_map_ifneeded()

            if (type(bin) != type('hello')):

                raise 'Error: Bin must be specified as a string.'

            if (os.path.exists(bin) == False):

                raise 'Error: File does not exist.'
            if rawcal:
                return(self.load_kv_from_file(filename=bin, calbase='rawcaldict'))
            return(self.load_kv_from_file(filename=bin))
                       
    # ---- Temporary EVB test command ----            
    def dacEval(self, ndx=0, fval = 'NAN'):
        if fval=='NAN':
            modify = 0
        else:
            modify = 1
        newdata = self.reg_float(regnum=0xCE, val=fval, index=1, index2=ndx, write=modify, option=0, remote=0, units=None)
#        newdata = self.reg(regnum=0xCE, index=ndx, index2=0x11, write=0)
        return(newdata)

    # ---- Temporary EVB test command ----            
    def dacEvalReadBack(self, ndx=0):
        newdata = self.reg_float(regnum=0xCE, val=0, index=2, index2=ndx, write=0, option=0, remote=0, units=None)
        return(newdata)

    # ---- Temporary EVB test command ----            
    def adcEval(self, ndx=0):
        newdata = self.reg_float(regnum=0xCE, val=0, index=0, index2=ndx, write=0, option=0, remote=0, units=None)
#        newdata = self.reg(regnum=0xCE, index=ndx, index2=0x11, write=0)
        return(newdata)

    # ---- Temporary EVB test command ----            
    def IOEval(self,ndx=0, pinVal='None'):
        if pinVal=='None':
            modify = 0
            pinVal = 0
        else:
            modify = 1
        newdata = self.reg(regnum=0xCF, index=ndx, index2=pinVal, write=modify)
#        newdata = self.reg(regnum=0xCE, index=ndx, index2=0x11, write=0)
        return(newdata)
                            